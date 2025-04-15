from docx import Document

FILENAME = "./data/TestFile.docx"

f = open(FILENAME,'rb')
doc = Document(f)

def iter_readings(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            yield paragraph

def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph.text

def parse_table(table):
    for row_idx,row in enumerate(table.rows):
        for col_idx,col in enumerate(row.cells):
            if col.tables:
                for nested_table in col.tables:
                    parse_table(nested_table)
            else:
                yield col.text

def iter_sections(doc):
    for part in doc.iter_inner_content():
        if part.style.name.startswith('Heading'):
            yield "hd: "+part.text
        elif "table" in part.style.name.lower():
            yield "tbl: " + f"{[val for val in parse_table(part)]}"
        else:
            yield part.text


if __name__ == "__main__":
    for heading in iter_sections(doc):
        print(heading)

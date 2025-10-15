from docx import Document as DocParser
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
import re
from collections.abc import Callable
from utils import getFirstPageOfDocxInMarkdown, getFirstTwoPagesOfDocxInMarkdown, getMetadataFromLLM, getCRContentFromLLM,Document
from concurrent.futures import ProcessPoolExecutor, as_completed

BASE_SECTION_NAME = "N/A"

def parse_table(table):
    """@input: docx table. Consists of rows and cells and the cells may have other tables.
    This is a helper function to help parse tables from docx. 
    It is needed as tables may contain nested tables."""
    for row_idx,row in enumerate(table.rows):
        for col_idx,col in enumerate(row.cells):
            if col.tables:
                for nested_table in col.tables:
                    parse_table(nested_table)
            else:
                yield col.text

def clean_file_name(name:str):
    """@name: the full name of the file with the path.
    This function returns just the filename without the path"""
    return name.split("/")[-1]

def process_section_name(section_name:str)->str:
    """@input: (str)  section name.
    This is a helper function to process the section name into a numerical format.
    We do this to keep it consistent with how external references will be searched for.
    """
    return section_name.split("\t")[0]

def extract_core_properties(file)->dict:
    """This is a helper to pull out the core properties of a docx file"""
    core_obj = file.core_properties
    return {'author':core_obj.author,'title':core_obj.title,'subject':core_obj.subject}

def addExtraDocumentWideMetadataForContext(text_chunk:str,filepath:str):
    """Use this for 3gpp specs"""
    extractVersionAndDocIDRegx = re.compile(r"(3GPP TS (\d+.\d+|\-\d)+ V\d+.\d+.\d)",re.IGNORECASE)
    extractVersion = re.compile(r"(V\d+.\d+.\d)")
    extractDocument = re.compile(r"(TS (\d+.\d+|\-\d)+)")
    extractTimestamp = re.compile(r"\((\d{4}-\d{2})\)")

    searchRes = extractVersionAndDocIDRegx.search(text_chunk)
    timestamp = extractTimestamp.search(text_chunk)

    if searchRes and timestamp:
        res = searchRes.group()
        version = extractVersion.search(res).group()[1:]
        docID = extractDocument.search(res).group()[3:]
        timestamp = timestamp.group().replace("(","").replace(")","")
    else:
        return {}
    
    metadata = {'version':version,'docID':docID,'timestamp':timestamp}
    return metadata

def addExtraDocumentWideMetadataForReason(text_chunk:str,filepath:str):
    """Use this for TDocs"""
    mdData = getFirstPageOfDocxInMarkdown(filepath)
    metadata =getMetadataFromLLM(mdData)
    print(metadata)
    return metadata

def section_chunks_of_file(file:str, addExtraDocumentWideMetadata:Callable[[str,str],dict]):
    f = open(file,'rb')
    try:
        doc = DocParser(f)
    except:
        raise Exception(f"for document {file} cannot parse with Docx")

    file_metadata = extract_core_properties(doc)

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000,chunk_overlap=200,add_start_index=True)

    current_section_title = BASE_SECTION_NAME
    current_section_text = ""
    sections = []
    #TODO: Switch to collecting a "current section" list of text chunks
    # currently there's no overlap between paragraphs in the same section which is bad for retrieval
    # We want to collect (section_text,current_section) where section_text is a combination of paragraphs
    # Then, we use the text_splitter to split it into split chunks
    for part in doc.iter_inner_content():
        #print(f"part is {part}")
        if isinstance(part,docx.text.paragraph.Paragraph) and part.style and part.style.name.startswith('Heading'):
            # update current section
            sections.append((current_section_text,current_section_title))
            current_section_text = ""
            current_section_title = part.text
        elif isinstance(part,docx.table.Table):
            for table_datum in parse_table(part):
                current_section_text += table_datum
        else:
            # append text to current section
            current_section_text += part.text
    
    #Add last section to sections
    sections.append((current_section_text,current_section_title))

    chunks_with_metadata = []
    addMetadata = {}

    for text,section_name in sections:
        split_chunks = text_splitter.split_text(text)
        
        if section_name == BASE_SECTION_NAME and not addMetadata:
            addMetadata = addExtraDocumentWideMetadata(text,file)
            print(f"metadata is {addMetadata}")

        for chunk in split_chunks:
            chunks_with_metadata.append(Document(
                page_content=chunk,
                metadata={'source':clean_file_name(file),'section':process_section_name(section_name),**addMetadata,**file_metadata}
            ))
    return chunks_with_metadata

def getSectionedChunks(file_list,addExtraDocumentWideMetadata:Callable[[str,str],dict]=addExtraDocumentWideMetadataForContext):
    """@input: file_list. List of files in relative path that will be chunked.
    @addExtraDocumentWideMetadata: func that returns a dictionary with extra metadata that will be added to all chunks.
    Returns: master list chunks_with_metadata that has chunks of all the files stored as Documents.
    These Documents have section metadata"""
    if len(file_list) == 1:
        return section_chunks_of_file(file_list[0], addExtraDocumentWideMetadata)
    else:
        chunks_with_metadata= []
        with ProcessPoolExecutor() as executor:
            futures = {executor.submit(section_chunks_of_file,file,addExtraDocumentWideMetadata): file for file in file_list}
            for future in as_completed(futures):
                chunks_with_metadata.extend(future.result())
    return chunks_with_metadata

def section_entire_chunks_of_file(file:str, addExtraDocumentWideMetadata:Callable[[str,str],dict]):
    f = open(file,'rb')
    try:
        doc = DocParser(f)
    except:
        raise Exception(f"for document {file} cannot parse with Docx")

    file_metadata = extract_core_properties(doc)

    current_section_title = BASE_SECTION_NAME
    current_section_text = ""
    sections = []

    for part in doc.iter_inner_content():
        if isinstance(part,docx.text.paragraph.Paragraph) and part.style and part.style.name.startswith('Heading'):
            # update current section
            if current_section_text.strip() != '':
                sections.append((current_section_text,current_section_title))
            current_section_text = ""
            current_section_title = part.text
        elif isinstance(part,docx.table.Table):
            for table_datum in parse_table(part):
                current_section_text += table_datum
        else:
            # append text to current section
            current_section_text += part.text
        
        #Add last section to sections
    if current_section_text.strip() != '': # avoid adding empty sections
        sections.append((current_section_text,current_section_title))

    chunks_with_metadata = []
    addMetadata = {}
    for text,section_name in sections:
        if section_name == BASE_SECTION_NAME and not addMetadata:
            addMetadata = addExtraDocumentWideMetadata(text,file)
            print(f"metadata is {addMetadata}")
        if section_name.strip() != '' and section_name.strip() != '–' and section_name.strip() != '-':
            chunks_with_metadata.append(Document(
                page_content=text,
                metadata = {'source':clean_file_name(file),'section':process_section_name(section_name),**addMetadata,**file_metadata}
            ))
    return chunks_with_metadata

def getFullSectionChunks(file_list,addExtraDocumentWideMetadata:Callable[[str,str],dict]=addExtraDocumentWideMetadataForContext):
    """@input: file_list. List of files in relative path that will be chunked.
    @addExtraDocumentWideMetadata: func that returns a dictionary with extra metadata that will be added to all chunks.
    this function creates chunks which are the size of an entire section rather than breaking the section up."""
    if len(file_list) == 1:
        return section_entire_chunks_of_file(file_list[0], addExtraDocumentWideMetadata)
    else:
        chunks_with_metadata= []
        with ProcessPoolExecutor() as executor:
            futures = {executor.submit(section_entire_chunks_of_file,file,addExtraDocumentWideMetadata): file for file in file_list}
            for future in as_completed(futures):
                chunks_with_metadata.extend(future.result())
    return chunks_with_metadata

def getFullFileChunks(file_list):
    chunks = []
    for file in file_list:

        f = open(file,'rb')
        doc = DocParser(f)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000,chunk_overlap=200,add_start_index=True)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text
        split_chunks = text_splitter.split_text(text)
        for chunk in split_chunks:
            chunks.append(Document(
                page_content=chunk,
                metadata={'source':clean_file_name(file)}
            ))
    return chunks


#####################
#### CR chunking ####
#####################
def process_cr_file(file:str):
    f = open(file,'rb')
    try:
        doc = DocParser(f)
    except:
        raise Exception(f"for document {file} cannot parse with Docx")

    file_metadata = extract_core_properties(doc)
    mdContent = getFirstTwoPagesOfDocxInMarkdown(file)
    metadata_from_llm = getMetadataFromLLM(mdContent)
    all_metadata = {**file_metadata,**metadata_from_llm}
    print(f"metadata is {all_metadata}")

    change_chunks: list[dict] = getCRContentFromLLM(mdContent)
    chunks_with_metadata = []
    for change_chunk in change_chunks:
        page_content = f"Change Summary: {change_chunk.get('summary','N/A')}\nReason: {change_chunk.get('reason','N/A')}\nConsequence if this change is not accepted: {change_chunk.get('consequence','N/A')}"
        chunks_with_metadata.append(Document(
            page_content=page_content,
            metadata={'source':clean_file_name(file),**all_metadata}
        ))
    return chunks_with_metadata

def getCRChunks(file_list:list[str])->list[Document]:
    if len(file_list) == 1:
        return process_cr_file(file_list[0])
    else:
        chunks = []
        with ProcessPoolExecutor() as executor:
            futures = {executor.submit(process_cr_file,file): file for file in file_list}
            for future in as_completed(futures):
                chunks.extend(future.result())
        return chunks

if __name__ =="__main__":
    """file_list = ["./data/38214-hc0.docx","./data/v16diffver.docx"]
    #print(getSectionedChunks(file_list)[80:90])
    sectioned_things = {}
    for doc in getFullSectionChunks([file_list[0]]):
        sectioned_things[doc.metadata["section"]] = sectioned_things.get(doc.metadata["section"],[]) + [doc.page_content]

    for doc in getFullSectionChunks([file_list[1]]):
        sectioned_things[doc.metadata["section"]] = sectioned_things.get(doc.metadata["section"],[]) + [doc.page_content]
    
    print(sectioned_things['2'])"""
    #convertAllDocToDocx("./reasoning")
    file_list = ["./all3gppdocsfromrel17and18/docsfordiffs/38331-h30.docx"]
    #print(Docx2txtLoader(file_list[0]).load())
    chunks = getFullSectionChunks(file_list,addExtraDocumentWideMetadataForReason)
    print(chunks[:100])
    seen = set()
    sectionToChunks = {}
    for chunk in chunks:
        if chunk.metadata["section"] in seen:
            print(f"duplicate section {chunk.metadata['section']},")
        else:
            seen.add(chunk.metadata["section"])
        sectionToChunks[chunk.metadata["section"]] = sectionToChunks.get(chunk.metadata["section"],[]) + [chunk.page_content]
    
    